import docx
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
import os

'''
Taking the files from the local directory. A ML library sumy to process larger document text into the character limits for prefered meta description length. 
'''

def process_all_docx_files(directory):
    # List all files in the given directory
    absolute_directory_path = os.path.abspath(directory)
    print(f"Processing files in directory: {absolute_directory_path}")
    # Iterate over all files in the directory
    for filename in os.listdir(directory):
        # Check if the file is a .docx file
        if filename.endswith(".docx"):
            file_path = os.path.join(directory, filename)
            try:
            # Check if the file is accessible
                if os.path.isfile(file_path) and os.access(file_path, os.R_OK):
                    print(f"Processing {filename}...")

                    # Extract title, version, and text from the document
                    title = extract_title_from_docx(file_path)
                    version = extract_version_from_docx(file_path)
                    policy_text = extract_text_from_docx(file_path)

                    # Summarize the text
                    summary = summarize_text(policy_text)

                    # Print or save the results
                    print(filename)
                    print(f"Title: {title}")
                    print(f"Version: {version}")
                    print(f"Summary: {summary}")
                    print("\n" + "#" * 80 + "\n")

                else:
                    print(f"Error: {filename} is not accessible.")
            except Exception as e:
                print(f"An error occurred while processing {filename}: {e}")

#Extract full text from .docx
def extract_text_from_docx(file_path):
    doc = docx.Document(file_path) # Open the .docx file
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text) # Add each paragraph's text to the list
    return '\n'.join(full_text) # Join all paragraphs with newline characters

def extract_title_from_docx(file_path):
    # Logic to extract title from Word document
    doc = docx.Document(file_path)
    # Extracting the first line of the document as the title
    title = doc.paragraphs[0].text
    # print(f"Extracted Title from DOCX: {title}")
    return title.strip()


def extract_version_from_docx(file_path):
    doc = docx.Document(file_path)
    version = None

    for para in doc.paragraphs:
        if "Policy Number:" in para.text:
            # Extracting the version part from the policy number line
            try:
                version = para.text.split('Policy Number:')[1].strip()
                break  # Exit the loop once the version is found
            except IndexError:
                # Handle the case where the line doesn't have the expected format
                pass
    # print(f"Extracted Title from DOCX: {version}")
    return version


def summarize_text(text, sentences_count=2):
    parser = PlaintextParser.from_string(text, Tokenizer("english"))
    summarizer = LsaSummarizer()
    summary = summarizer(parser.document, sentences_count) # Summarize the document

    return " ".join([str(sentence) for sentence in summary])  # Return the summary as a string


"""
File directory and setting variables
"""
# directory_path = "documents"
# process_all_docx_files(directory_path)
